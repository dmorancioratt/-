import io
import re
import zipfile
from pathlib import Path
from typing import BinaryIO

try:
    from docx import Document
except ImportError:  # pragma: no cover - python-docx 未安装时降级，登录/非 docx 功能不受影响
    Document = None  # type: ignore[assignment,misc]

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - pypdf 未安装时降级，登录/非 pdf 功能不受影响
    PdfReader = None  # type: ignore[assignment,misc]


MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 120_000
MAX_PDF_PAGES = 80
MAX_DOCX_UNCOMPRESSED_BYTES = 60 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class DocumentParseError(ValueError):
    def __init__(self, message: str, *, status_code: int = 422):
        super().__init__(message)
        self.status_code = status_code


def extract_resume_text(filename: str, content: bytes) -> tuple[str, str]:
    safe_name = Path(filename or "resume").name
    extension = Path(safe_name).suffix.lower()
    if extension == ".doc":
        raise DocumentParseError("暂不支持旧版 .doc 文件，请在 Word 中另存为 .docx 后重新上传", status_code=415)
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError("仅支持 PDF、DOCX、TXT 和 Markdown 简历文件", status_code=415)
    if not content:
        raise DocumentParseError("上传的简历文件为空")
    if len(content) > MAX_UPLOAD_BYTES:
        raise DocumentParseError("简历文件不能超过 10MB", status_code=413)

    if extension == ".pdf":
        text = _extract_pdf(content)
        file_type = "pdf"
    elif extension == ".docx":
        text = _extract_docx(content)
        file_type = "docx"
    else:
        text = _extract_plain_text(content)
        file_type = "text"

    normalized = _normalize_text(text)
    if len(normalized) < 20:
        if extension == ".pdf":
            raise DocumentParseError("PDF 中没有足够的可提取文字；如果这是扫描版或图片简历，请先进行 OCR 后再上传")
        raise DocumentParseError("文件中没有足够的简历文字，请检查文件内容")
    if len(normalized) > MAX_EXTRACTED_CHARACTERS:
        normalized = normalized[:MAX_EXTRACTED_CHARACTERS]
    return normalized, file_type


def _extract_pdf(content: bytes) -> str:
    if PdfReader is None:
        raise DocumentParseError("后端未安装 pypdf，请先执行 pip install pypdf 后重启服务", status_code=503)
    if not content.startswith(b"%PDF"):
        raise DocumentParseError("文件扩展名是 PDF，但文件内容不是有效 PDF")
    try:
        reader = PdfReader(io.BytesIO(content), strict=False)
        if reader.is_encrypted:
            try:
                unlocked = reader.decrypt("")
            except Exception as exc:
                raise DocumentParseError("PDF 已加密，请解除密码保护后重新上传") from exc
            if not unlocked:
                raise DocumentParseError("PDF 已加密，请解除密码保护后重新上传")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise DocumentParseError(f"PDF 页数不能超过 {MAX_PDF_PAGES} 页", status_code=413)
        pages = []
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text(extraction_mode="layout") or page.extract_text() or ""
            except Exception:
                page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"[第 {page_number} 页]\n{page_text.strip()}")
        return "\n\n".join(pages)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError("PDF 文件损坏或格式不受支持，无法读取") from exc


def _extract_docx(content: bytes) -> str:
    if Document is None:
        raise DocumentParseError("后端未安装 python-docx，请先执行 pip install python-docx 后重启服务", status_code=503)
    if not content.startswith(b"PK"):
        raise DocumentParseError("文件扩展名是 DOCX，但文件内容不是有效 Word 文档")
    _validate_docx_archive(content)
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise DocumentParseError("Word 文档损坏或格式不受支持，无法读取") from exc

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        _append_unique(blocks, paragraph.text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            _append_unique(blocks, row_text)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            _append_unique(blocks, paragraph.text)
        for paragraph in section.footer.paragraphs:
            _append_unique(blocks, paragraph.text)

    primary_text = "\n".join(blocks)
    ooxml_text = _extract_ooxml_text(document)
    return ooxml_text if len(ooxml_text) > len(primary_text) * 1.15 else primary_text


def _validate_docx_archive(content: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise DocumentParseError("该文件不是有效的 DOCX Word 文档")
            total_size = sum(info.file_size for info in archive.infolist())
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise DocumentParseError("Word 文档解压后过大，无法安全处理", status_code=413)
    except DocumentParseError:
        raise
    except zipfile.BadZipFile as exc:
        raise DocumentParseError("Word 文档损坏，无法解压读取") from exc


def _extract_ooxml_text(document: Document) -> str:
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs = []
    for paragraph in document.element.body.iter(f"{namespace}p"):
        text = "".join(node.text or "" for node in paragraph.iter(f"{namespace}t")).strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


def _extract_plain_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("文本文件编码无法识别，请转换为 UTF-8 后重新上传")


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _append_unique(items: list[str], value: str) -> None:
    normalized = _normalize_text(value)
    if normalized and normalized not in items:
        items.append(normalized)
