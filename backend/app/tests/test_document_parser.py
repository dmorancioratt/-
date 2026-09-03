from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.main import app
from app.services.document_parser import DocumentParseError, extract_resume_text


client = TestClient(app)


def build_docx() -> bytes:
    document = Document()
    document.add_paragraph("姓名：测试候选人")
    document.add_paragraph("本科计算机科学，熟悉 Python、SQL 和数据分析。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目经历"
    table.cell(0, 1).text = "招聘数据分析平台"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf(text: str = "Resume Python SQL project experience and data analysis") -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def auth_headers() -> dict[str, str]:
    login = client.post("/api/auth/login", json={"username": "student_demo", "password": "Demo@123"})
    assert login.status_code == 200
    return {"Authorization": f"Bearer {login.json()['token']}"}


def test_extracts_pdf_text():
    text, file_type = extract_resume_text("resume.pdf", build_pdf())
    assert file_type == "pdf"
    assert "Python SQL" in text


def test_extracts_docx_paragraphs_and_tables():
    text, file_type = extract_resume_text("resume.docx", build_docx())
    assert file_type == "docx"
    assert "测试候选人" in text
    assert "招聘数据分析平台" in text


@pytest.mark.parametrize(
    ("filename", "content", "content_type"),
    [
        ("resume.pdf", build_pdf(), "application/pdf"),
        ("resume.docx", build_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    ],
    ids=["pdf", "docx"],
)
def test_resume_file_endpoint(filename: str, content: bytes, content_type: str):
    headers = auth_headers()
    response = client.post(
        "/api/resume/parse-file",
        headers=headers,
        files={"file": (filename, content, content_type)},
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["file"]["character_count"] >= 20
    assert payload["result"]["ai_provider"] == "mock"
    assert payload["result"]["resume_id"] > 0
    assert payload["extracted_text"]
    history = client.get("/api/resumes", headers=headers)
    assert history.status_code == 200
    saved = next(row for row in history.json() if row["id"] == payload["result"]["resume_id"])
    assert saved["source_filename"] == filename
    assert saved["created_at"]


def test_rejects_scanned_pdf_without_text():
    with pytest.raises(DocumentParseError, match="OCR"):
        extract_resume_text("scan.pdf", build_pdf(""))


def test_rejects_legacy_doc():
    with pytest.raises(DocumentParseError, match=".docx"):
        extract_resume_text("resume.doc", b"legacy word file")
