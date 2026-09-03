from copy import deepcopy
from pathlib import Path
import re
import shutil

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\wwcc1\Desktop\挑战杯ai生成版文档以及要求\最终成果-数融智联")
SOURCE = ROOT / "_过程备份" / "20260901-会议重构前" / "数融智联-挑战杯项目书-48页版.docx"
OUTPUT = ROOT / "材料文档" / "数融智联-挑战杯项目书-修改版.docx"
REFERENCE = ROOT / "_过程备份" / "new-meeting-reference"
ASSETS = ROOT / "图片素材"
FIGURES = ASSETS / "image2科研重绘版"
LOGO = ASSETS / "挑战杯通用标识.png"
COVER = ASSETS / "封面科研背景-会议优化版.png"


def set_run_font(run, size=10.5, bold=False, color="303846", name="宋体"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_heading(paragraph, text, level):
    paragraph.clear()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.space_before = Pt(8 if level > 1 else 0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, 16 if level == 1 else 13 if level == 2 else 11, True, "0B3A66")


def element_paragraph(element, doc):
    return Paragraph(element, doc._body)


def paragraph_text(element, doc):
    if element.tag != qn("w:p"):
        return ""
    return element_paragraph(element, doc).text.strip()


def paragraph_style(element, doc):
    if element.tag != qn("w:p"):
        return ""
    return element_paragraph(element, doc).style.name


def create_cover_assets():
    shutil.copy2(REFERENCE / "logo-candidate-1.png", LOGO)

    width, height = 2480, 3508
    image = Image.new("RGB", (width, height), "#F7FAFC")
    draw = ImageDraw.Draw(image, "RGBA")

    draw.rectangle((0, 0, width, 180), fill="#0B3A66")
    draw.rectangle((0, 180, width, 206), fill="#18A6A6")
    for x in range(80, width, 120):
        draw.line((x, 206, x, height), fill=(44, 86, 120, 13), width=2)
    for y in range(260, height, 120):
        draw.line((0, y, width, y), fill=(44, 86, 120, 13), width=2)

    nodes = [
        (1600, 510), (1830, 650), (2080, 540), (2250, 780), (1990, 920),
        (2180, 1120), (1880, 1230), (2280, 1430), (2050, 1610), (2350, 1840),
        (1800, 1940), (2140, 2180), (2340, 2440), (1910, 2570), (2210, 2820),
    ]
    for i in range(len(nodes) - 1):
        draw.line((nodes[i], nodes[i + 1]), fill=(11, 58, 102, 40), width=5)
    for i in range(0, len(nodes) - 3, 3):
        draw.line((nodes[i], nodes[i + 3]), fill=(24, 166, 166, 42), width=4)
    for x, y in nodes:
        draw.ellipse((x - 17, y - 17, x + 17, y + 17), fill=(24, 166, 166, 105), outline=(11, 58, 102, 120), width=3)

    draw.polygon([(0, 2760), (1130, 3508), (0, 3508)], fill=(11, 58, 102, 24))
    draw.polygon([(0, 3060), (670, 3508), (0, 3508)], fill=(24, 166, 166, 32))
    draw.line((260, 2820, 1430, 2820), fill=(11, 58, 102, 52), width=4)
    draw.line((260, 2860, 1110, 2860), fill=(24, 166, 166, 52), width=4)

    logo = Image.open(LOGO).convert("RGBA")
    pixels = logo.load()
    for y in range(logo.height):
        for x in range(logo.width):
            r, g, b, a = pixels[x, y]
            if r > 245 and g > 245 and b > 245:
                pixels[x, y] = (255, 255, 255, 0)
    bbox = logo.getbbox()
    logo = logo.crop(bbox).resize((430, 430), Image.Resampling.LANCZOS)
    image.paste(logo, ((width - logo.width) // 2, 260), logo)
    image.save(COVER, quality=96, dpi=(300, 300))


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_heading(doc, text, level=1, page_break=False):
    if page_break:
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph(style=f"Heading {level}")
    set_heading(p, text, level)
    return p._p


def add_body(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_run_font(p.add_run(text))
    return p._p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "DCEAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(header), 9, True, "0B3A66")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(str(value)), 8.6)
    if widths:
        for row in table.rows:
            for index, value in enumerate(widths):
                row.cells[index].width = Cm(value)
    return table._tbl


def add_figure(doc, filename, caption):
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_with_next = True
    picture.add_run().add_picture(str(FIGURES / filename), width=Cm(16.2))
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_before = Pt(3)
    label.paragraph_format.space_after = Pt(6)
    set_run_font(label.add_run(caption), 9, False, "4B5563")
    return [picture._p, label._p]


def add_link_directory(doc):
    elements = []
    elements.append(add_heading(doc, "演示与提交链接目录", 1, page_break=True))
    elements.append(add_body(
        doc,
        "本页集中放置评审人员需要访问的系统入口与提交凭据。本地入口已经过复测；公网演示、云盘地址、提取码和上传时间截图须在实际上传后由团队核对，未核实信息不得用示例值替代。",
    ))
    elements.append(add_heading(doc, "一、系统访问入口", 2))
    elements.append(add_table(doc, ["项目", "地址或状态", "说明"], [
        ["前端本地入口", "http://127.0.0.1:5173", "执行启动脚本后访问"],
        ["后端接口文档", "http://127.0.0.1:8000/docs", "用于接口检查与复测"],
        ["公网演示地址", "待确认", "部署完成后填写最终地址"],
        ["演示视频", "待确认", "录制完成后填写文件名或链接"],
    ], [3.6, 6.5, 6.0]))
    elements.append(add_heading(doc, "二、云盘提交信息", 2))
    elements.append(add_table(doc, ["材料", "填写内容", "提交前检查"], [
        ["安全云盘链接", "待确认", "链接可访问且有效期覆盖评审期"],
        ["提取码", "待确认", "与链接逐字核对"],
        ["上传时间", "待确认", "以云盘页面显示时间为准"],
        ["上传时间截图", "待确认：上传后替换本行", "截图需同时显示文件名与时间"],
    ], [4.0, 6.0, 6.1]))
    elements.append(add_heading(doc, "三、演示账号", 2))
    elements.append(add_table(doc, ["角色", "账号", "密码", "主要入口"], [
        ["平台管理员", "admin_demo", "Demo@123", "数据源、审核、评测"],
        ["企业HR", "hr_admin", "Demo@123", "岗位、JD解析、候选人"],
        ["求职者", "student_demo", "Demo@123", "画像、匹配、学习路径"],
        ["候选人样例", "candidate_demo", "Demo@123", "简历解析与成长反馈"],
    ], [3.0, 4.0, 3.5, 5.6]))
    elements.append(add_heading(doc, "四、佐证材料", 2))
    elements.append(add_table(doc, ["佐证项", "当前状态", "提交要求"], [
        ["软件著作权证书", "待确认", "证书原件尚未提供；取得后仅附证书页"],
        ["系统测试报告", "已具备", "测试数据目录及 evaluation_summary.json"],
        ["部署说明", "已具备", "材料文档/数融智联-作品部署说明.docx"],
    ], [4.0, 3.2, 8.9]))
    return elements


def add_requirement_matrix(doc):
    return [
        add_heading(doc, "2.5 功能与非功能需求", 2),
        add_body(doc, "系统需求按可用、可信、可部署三个层次收敛。功能需求对应评委可操作的完整流程，非功能需求约束数据边界、解释方式和部署条件，验收口径直接落到页面、接口、日志或测试报告。"),
        add_table(doc, ["需求类别", "具体要求", "验收证据"], [
            ["数据治理", "支持来源登记、授权记录、哈希去重、字段证据与人工复核", "数据源页面、审核任务、证据抽屉"],
            ["岗位图谱", "保存岗位、技能、证书、关系、版本与演化事件", "图谱探索、能力演化、版本记录"],
            ["智能分析", "JD/简历解析、RAG校验、候选新岗位、匹配解释", "解析结果、错误案例、评分明细"],
            ["用户闭环", "HR筛选、求职者画像、学习路径、管理员治理", "三类账号完整演示流程"],
            ["安全与隐私", "最小化采集、角色隔离、敏感字段不进入外部模型", "权限测试、环境变量、审计记录"],
            ["可部署性", "Windows一键启动、前后端分离、Docker Compose可选", "部署说明与本地复现记录"],
        ], [3.0, 8.0, 5.1]),
    ]


def add_module_overview(doc):
    return [
        add_body(doc, "核心模块不是彼此孤立的功能页。数据治理模块提供可追溯事实，图谱模块负责结构化与版本化，智能分析模块在证据边界内完成抽取和解释，人工协同模块处理冲突与低置信度结果，可视化模块把结论还原为评委能够检查的来源、关系和版本。"),
        add_table(doc, ["模块", "输入", "关键处理", "输出与人工控制点"], [
            ["多源数据治理", "权威来源、授权JD", "清洗、去重、字段溯源、置信度", "标准化事实；冲突进入复核"],
            ["图谱构建与演化", "岗位、技能、证书、证据", "本体映射、关系入库、版本差异", "岗位画像、演化事件；发布前审核"],
            ["RAG双层校验", "候选字段、检索证据", "检索约束、图谱事实校验", "带证据解释；无依据则拒答"],
            ["解析与匹配", "JD、授权简历、岗位画像", "结构化解析、六维确定性评分", "评分明细、缺口；分值不可由模型改写"],
            ["人机协同", "冲突、低置信度、用户反馈", "编辑、驳回、回流与版本记录", "正式数据或错误样本"],
            ["可视化交互", "图谱、演化、评测结果", "多视图联动与证据下钻", "可检查页面与导出结果"],
        ], [3.2, 3.5, 5.0, 4.4]),
    ]


def add_ablation_plan(doc):
    return [
        add_heading(doc, "7.9 消融实验与创新模块增益验证", 2),
        add_body(doc, "现有14条金标样本只承担回归检查，不能据此宣称创新模块具有稳定增益。消融实验将固定120条模拟JD的数据版本和评测脚本，逐项关闭多源权重、时间衰减、图谱事实校验与技能共现特征，比较候选误报、字段抽取和Top-1变化。正式数字以同一环境复测结果为准。"),
        add_table(doc, ["实验组", "关闭或替换模块", "主要观察指标", "当前状态"], [
            ["完整系统", "无", "JD F1、匹配Top-1、候选误报率", "回归指标已复现"],
            ["Ablation-A", "取消来源权威度与多源一致性", "错误来源进入正式事实的比例", "待复测"],
            ["Ablation-B", "取消时间衰减", "过时技能权重与候选排序变化", "待复测"],
            ["Ablation-C", "取消图谱事实二次校验", "无证据字段与幻觉拦截率", "待复测"],
            ["Ablation-D", "技能共现改为关键词频次", "新岗位候选Precision/Recall", "待补充金标后复测"],
        ], [3.0, 5.6, 5.0, 2.5]),
    ]


def add_implementation_inventory(doc):
    return [
        add_heading(doc, "6.6 核心代码模块与启动链路", 2),
        add_body(doc, "源码按前端、后端、数据和部署脚本分开组织。业务规则集中在后端服务层，接口层只负责参数校验和响应封装；前端页面通过统一API访问数据，不在浏览器中保存数据库口令或模型密钥。这样的拆分便于评审时定位功能，也便于后续替换数据库或AI服务。"),
        add_table(doc, ["目录或文件", "功能", "评审时可检查内容"], [
            ["backend/app/routers", "认证、岗位、图谱、评测与工作流接口", "接口路径、角色守卫、错误响应"],
            ["backend/app/services", "数据治理、解析、RAG、匹配与幻觉防控", "规则参数、证据校验、人工复核入口"],
            ["backend/app/models.py", "岗位、技能、证书、版本与审核任务模型", "字段定义、关系约束、审计字段"],
            ["backend/app/evaluation", "金标读取、预测生成和离线评测", "指标公式、错误案例、JSON报告"],
            ["frontend/src/views", "三类用户页面与演示流程", "数据源、图谱、演化、匹配、评测页面"],
            ["start.bat / stop.bat", "Windows一键启动与停止", "端口检查、进程清理、日志位置"],
            ["Dockerfile / compose", "容器构建与前后端编排", "依赖安装、端口、数据卷与启动命令"],
        ], [4.2, 5.6, 6.0]),
        add_body(doc, "本地启动顺序为数据库初始化、FastAPI服务、Vue前端。测试环境固定设置 APP_ENV=test 和 AI_PROVIDER=mock，不访问外部模型；演示环境若启用DeepSeek，只由后端读取密钥。任何包含密钥、缓存、日志或未授权简历的文件都不进入提交压缩包。"),
    ]


def add_test_case_matrix(doc):
    return [
        add_heading(doc, "7.7 测试样本分组与异常场景", 2),
        add_body(doc, "120条模拟JD用于检查导入、解析、去重、图谱更新和审核链路，不作为真实市场样本。样本按正常岗位、重复模板、过时技能、弱岗位标题、字段缺失和跨领域技能组合分组；每条样本保存预期处理状态，测试失败时能够回到具体输入。"),
        add_table(doc, ["样本组", "主要扰动", "预期行为", "检查点"], [
            ["正常JD", "字段完整、来源清晰", "完成解析并生成字段证据", "必备/加分技能、证据片段"],
            ["重复模板JD", "正文复制、标题轻微变化", "按哈希和近似规则去重", "重复记录不放大多源一致性"],
            ["过时JD", "发布时间早、技能组合陈旧", "时间衰减后降低影响", "权重、版本和候选排序"],
            ["弱岗位标题", "萌芽岗、复合岗或营销式名称", "不直接认定新岗位", "技能共现与人工复核任务"],
            ["字段缺失", "职责或技能证据不足", "标记低置信度并进入复核", "拒绝无依据补全"],
            ["跨领域组合", "同名技能在不同行业出现", "结合场景与岗位层级消歧", "关系类型和证据来源"],
        ], [3.0, 4.0, 5.0, 4.0]),
        add_heading(doc, "7.8 性能、稳定性与图谱演化复测", 2),
        add_body(doc, "性能测试不使用一次页面加载时间代替系统结论。复测将固定评审电脑、数据库快照和并发参数，分别记录JD批量导入、图谱查询、匹配报告生成和评测脚本运行耗时；图谱演化测试则核对12条演化事件的新增、调整、淘汰和版本回滚。当前材料只列方法，正式数值待同一设备复测后填写。"),
        add_table(doc, ["测试对象", "记录指标", "通过口径", "状态"], [
            ["JD批量导入", "总耗时、失败数、重复拦截数", "120条全部形成明确处理状态", "流程已通过，耗时待复测"],
            ["图谱查询", "响应时间、空结果与错误率", "页面可返回并能下钻证据", "待统一设备复测"],
            ["匹配报告", "生成耗时、六维分值一致性", "同一输入重复运行分值一致", "确定性链路已验证"],
            ["能力演化", "事件正确率、版本回滚结果", "12条事件类型与预期一致", "待补充专项报告"],
            ["离线评测", "脚本耗时、JSON结果哈希", "同一版本重复运行结果一致", "已可复现"],
        ], [3.2, 4.8, 5.3, 3.0]),
    ]


def add_conclusion(doc):
    return [
        add_heading(doc, "8.10 总结与后续方向", 2),
        add_body(doc, "数融智联已经完成从来源登记、岗位事实治理、图谱建模、新岗位候选、能力演化到匹配解释的原型闭环。当前价值不在于用一个模型替代招聘判断，而在于把分散、易过时的岗位信息整理成可以核查、更新和复用的能力证据。"),
        add_body(doc, "后续工作集中在三件事：取得授权真实JD并扩充双人标注金标集；完成校内就业指导和企业岗位标准化试点；按统一数据版本执行消融实验与长期稳定性测试。试点单位、合作企业和软著证书在取得真实材料前均保持待确认状态。"),
    ]


def page_break_element(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p._p


def clone_range(children, doc, start_text, end_text=None):
    start = next(i for i, element in enumerate(children) if paragraph_text(element, doc) == start_text)
    if end_text:
        end = next(i for i, element in enumerate(children[start + 1:], start + 1) if paragraph_text(element, doc) == end_text)
    else:
        end = len(children) - 1
    return [deepcopy(element) for element in children[start:end]]


def clean_block(elements, doc, strip_figures=False):
    cleaned = []
    for element in elements:
        if element.tag == qn("w:p"):
            paragraph = element_paragraph(element, doc)
            for br in list(element.xpath(".//w:br")):
                if br.get(qn("w:type")) == "page":
                    br.getparent().remove(br)
            if strip_figures and element.xpath(".//w:drawing"):
                continue
            if strip_figures and re.match(r"^图\s*\d+", paragraph.text.strip()):
                continue
            if not paragraph.text.strip() and not element.xpath(".//w:drawing"):
                if cleaned and cleaned[-1].tag == qn("w:p") and not element_paragraph(cleaned[-1], doc).text.strip():
                    continue
        cleaned.append(element)
    return cleaned


def rewrite_block(elements, doc, mapping, strip_figures=False):
    elements = clean_block(elements, doc, strip_figures)
    for element in elements:
        if element.tag != qn("w:p"):
            continue
        paragraph = element_paragraph(element, doc)
        text = paragraph.text.strip()
        if text in mapping:
            new_text, level = mapping[text]
            set_heading(paragraph, new_text, level)
    return elements


def drop_last_table_with_caption(elements, doc):
    table_index = max(i for i, element in enumerate(elements) if element.tag == qn("w:tbl"))
    del elements[table_index]
    for index in range(table_index - 1, -1, -1):
        text = paragraph_text(elements[index], doc)
        if re.match(r"^表\s*\d+", text):
            del elements[index]
            break
        if text:
            break
    return elements


def generated(doc, producer):
    before = set(id(element) for element in doc.element.body)
    producer(doc)
    return [element for element in list(doc.element.body) if id(element) not in before and element.tag != qn("w:sectPr")]


def normalise_document(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.25)
        section.right_margin = Cm(2.05)
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Normal":
            paragraph.paragraph_format.orphan_control = True
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)


def renumber_captions(doc):
    figure_number = 0
    table_number = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^图\s*\d+", text):
            figure_number += 1
            rest = re.sub(r"^图\s*\d+\s*", "", text)
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run(f"图{figure_number} {rest}"), 9, False, "4B5563")
        elif re.match(r"^表\s*\d+", text):
            table_number += 1
            rest = re.sub(r"^表\s*\d+\s*", "", text)
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run(f"表{table_number} {rest}"), 9, False, "4B5563")
    return figure_number, table_number


def format_front(doc):
    first = doc.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.space_before = Pt(160)
    first.paragraph_format.space_after = Pt(28)
    first.clear()
    set_run_font(first.add_run("挑战杯项目申报材料"), 23, True, "0B3A66", "微软雅黑")

    second = doc.paragraphs[1]
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second.paragraph_format.space_after = Pt(180)
    second.clear()
    set_run_font(second.add_run("数融智联：多源异构数据驱动的岗位能力图谱构建与\n动态演化分析系统"), 20, True, "172033", "微软雅黑")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(30)
    set_run_font(info.add_run("申报单位：待确认    项目类别：待确认    2026年9月"), 11, False, "44546A")
    body = doc.element.body
    anchor = doc.paragraphs[2]._p
    body.remove(info._p)
    anchor.addprevious(info._p)

    paragraphs = doc.paragraphs
    toc_index = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "目录")
    toc_body = paragraphs[toc_index + 1]
    toc_body.clear()
    toc_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_body.paragraph_format.left_indent = Cm(2.0)
    toc_body.paragraph_format.line_spacing = 1.55
    toc_text = (
        "第一章  项目概述\n"
        "第二章  需求分析\n"
        "第三章  系统总体架构设计\n"
        "第四章  核心模块详细设计\n"
        "第五章  核心算法与创新点\n"
        "第六章  系统实现\n"
        "第七章  测试与评估\n"
        "第八章  总结与展望\n"
        "参考资料\n"
        "演示与提交链接目录"
    )
    set_run_font(toc_body.add_run(toc_text), 11.5, False, "26384A")

    paragraphs = doc.paragraphs
    abstract_index = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "摘要")
    abstract_body = paragraphs[abstract_index + 1:abstract_index + 5]
    abstract_texts = [
        "招聘岗位变化很快，但高校课程、职业标准和招聘JD的更新节奏并不一致。团队在整理人工智能相关岗位时发现，同一种能力常被不同名称描述，同一岗位名称在不同企业又可能对应不同职责。现有工具擅长检索和推荐，却很难回答一项能力来自哪里、何时发生变化、为什么影响匹配。数融智联因此把工作重点放在岗位事实治理：先保存来源与证据，再构建岗位能力图谱，最后开展新岗位发现、能力演化和匹配解释。",
        "系统由多源数据治理、岗位能力图谱、RAG双层校验、新岗位候选、六维确定性匹配和人工复核组成。前端采用Vue 3，后端采用FastAPI与SQLAlchemy，正式演示库使用SQLite，外部模型为可选能力。生成式模型可以协助抽取和组织语言，但不能绕过证据校验写入正式图谱，也不能修改匹配分值。",
        "截至2026年8月31日，正式演示库记录8个权威来源、107个岗位、140个技能、772条岗位—技能关系和127条岗位—证书关系；后端59项自动化测试全部通过。14条小规模金标样本上的JD抽取F1为93.88%、简历抽取F1为95.65%、匹配Top-1为80.00%。这些指标用于防止算法回归，不代表产业规模精度；真实招聘数据将在取得授权后接入。",
        "关键词：岗位能力图谱；多源异构数据；新岗位发现；能力动态演化；人岗匹配；可信AI",
    ]
    for paragraph, text in zip(abstract_body, abstract_texts):
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0.74) if not text.startswith("关键词") else Cm(0)
        paragraph.paragraph_format.line_spacing = 1.4
        set_run_font(paragraph.add_run(text), 10.5, False, "303846")


def main():
    create_cover_assets()
    doc = Document(SOURCE)
    original = list(doc.element.body)

    chapter1 = clone_range(original, doc, "第一章 项目背景与问题定义", "第二章 建设目标、功能需求与评审指标响应")
    goal = clone_range(original, doc, "2.1 评价维度与研究目标", "2.2 用户需求分层")
    users = clone_range(original, doc, "2.2 用户需求分层", "2.3 当前验证边界与待补证据")
    boundary = clone_range(original, doc, "2.3 当前验证边界与待补证据", "第三章 总体解决方案")
    solution = clone_range(original, doc, "第三章 总体解决方案", "第四章 数据基础与系统设计")
    chapter3 = clone_range(original, doc, "第四章 数据基础与系统设计", "第五章 核心技术路线与创新点")
    chapter5 = clone_range(original, doc, "第五章 核心技术路线与创新点", "第六章 系统实现与演示")
    chapter6 = clone_range(original, doc, "第六章 系统实现与演示", "第七章 测试验证与效果评估")
    chapter7 = clone_range(original, doc, "第七章 测试验证与效果评估", "第八章 应用场景、竞争分析与转化前景")
    chapter8 = clone_range(original, doc, "第八章 应用场景、竞争分析与转化前景", "第九章 团队执行计划与成果清单")
    references = clone_range(original, doc, "参考资料", "附录一 典型应用流程")

    appendices = {}
    appendix_titles = [
        "附录一 典型应用流程", "附录三 算法验证与可解释性说明", "附录八 数据来源分层、许可与使用边界",
        "附录九 多源数据治理与字段级证据链", "附录十 岗位能力图谱本体与关系模式",
        "附录十一 新岗位发现模型与判定规则", "附录十二 新岗位候选的人工复核准则",
        "附录十三 岗位能力动态演化与版本管理", "附录十四 六维确定性人岗匹配模型",
        "附录十五 匹配解释、学习路径与反馈更新", "附录十六 生成式模型的职责边界与幻觉防控",
        "附录十七 系统总体技术架构与模块依赖", "附录十八 三类用户协同流程与权限边界",
        "附录十九 原型系统页面与可演示证据", "附录二十 部署拓扑、运行方式与离线能力",
        "附录二十一 测试策略、数据基线与可复现性", "附录二十二 离线评测方法、指标定义与结果",
        "附录二十三 典型错误、原因定位与改进实验", "附录二十四 同类方案比较与项目创新定位",
        "附录二十五 应用场景、试点设计与验收指标", "附录二十六 社会价值与公共利益边界",
        "附录二十七 成果转化模式与成本结构", "附录二十八 技术成熟度、应用成熟度与推进路线",
        "附录三十 风险矩阵与应对预案",
    ]
    all_h1 = [paragraph_text(element, doc) for element in original if paragraph_style(element, doc) == "Heading 1"]
    for title in appendix_titles:
        pos = all_h1.index(title)
        end = all_h1[pos + 1] if pos + 1 < len(all_h1) else None
        appendices[title] = clone_range(original, doc, title, end)

    front_end = next(i for i, element in enumerate(original) if paragraph_text(element, doc) == "第一章 项目背景与问题定义")
    front = [deepcopy(element) for element in original[:front_end]]

    chapter1 = rewrite_block(chapter1, doc, {"第一章 项目背景与问题定义": ("第一章 项目概述", 1)})
    goal = rewrite_block(goal, doc, {"2.1 评价维度与研究目标": ("1.4 项目目标与评价指标", 2)})
    boundary = rewrite_block(boundary, doc, {"2.3 当前验证边界与待补证据": ("1.5 研究边界与证据口径", 2)})
    solution = rewrite_block(solution, doc, {
        "第三章 总体解决方案": ("1.6 总体思路与业务闭环", 2),
        "3.1 项目定位": ("1.6.1 项目定位", 3),
        "3.2 业务闭环": ("1.6.2 三类用户业务闭环", 3),
    })
    solution = drop_last_table_with_caption(solution, doc)

    users = rewrite_block(users, doc, {"2.2 用户需求分层": ("2.1 用户需求与痛点", 2)})
    app1 = rewrite_block(appendices["附录一 典型应用流程"], doc, {
        "附录一 典型应用流程": ("2.2 典型业务需求闭环", 2),
        "A1.1 企业端：从“写不清岗位”到“形成岗位标准”": ("2.2.1 企业端岗位标准化", 3),
        "A1.2 求职者端：从“泛泛投递”到“知道如何补齐差距”": ("2.2.2 求职者能力成长", 3),
        "A1.3 管理端：从“相信模型”到“审计模型输出”": ("2.2.3 管理端证据治理", 3),
    })
    app8 = rewrite_block(appendices["附录八 数据来源分层、许可与使用边界"], doc, {
        "附录八 数据来源分层、许可与使用边界": ("2.3 数据需求与授权边界", 2),
    })
    app24 = rewrite_block(appendices["附录二十四 同类方案比较与项目创新定位"], doc, {
        "附录二十四 同类方案比较与项目创新定位": ("2.4 同类方案与现有系统比较", 2),
    })

    chapter3 = rewrite_block(chapter3, doc, {
        "第四章 数据基础与系统设计": ("第三章 系统总体架构设计", 1),
        "4.1 数据来源与边界": ("3.1 数据层与来源边界", 2),
        "4.2 功能模块设计": ("3.2 功能模块划分", 2),
        "4.3 技术架构": ("3.3 全链路数据流与技术架构", 2),
    })
    app17 = rewrite_block(appendices["附录十七 系统总体技术架构与模块依赖"], doc, {
        "附录十七 系统总体技术架构与模块依赖": ("3.4 技术选型与模块依赖", 2),
    }, strip_figures=True)

    chapter4_parts = []
    chapter4_parts.extend(generated(doc, lambda d: add_heading(d, "第四章 核心模块详细设计", 1, page_break=True)))
    chapter4_parts.extend(generated(doc, lambda d: [d.element.body.append(e) for e in add_module_overview(d)]))
    module_specs = [
        ("附录九 多源数据治理与字段级证据链", "4.1 多源数据治理模块"),
        ("附录十 岗位能力图谱本体与关系模式", "4.2 图谱构建与演化模块"),
        ("附录十二 新岗位候选的人工复核准则", "4.3 人机协同复核模块"),
        ("附录十三 岗位能力动态演化与版本管理", "4.4 岗位能力演化与版本模块"),
        ("附录十五 匹配解释、学习路径与反馈更新", "4.5 简历解析、匹配与学习路径模块"),
        ("附录十六 生成式模型的职责边界与幻觉防控", "4.6 RAG双层校验与幻觉防控模块"),
        ("附录十八 三类用户协同流程与权限边界", "4.7 可视化交互与权限边界"),
    ]
    for source_title, new_title in module_specs:
        chapter4_parts.extend(rewrite_block(appendices[source_title], doc, {source_title: (new_title, 2)}, strip_figures=True))

    chapter5 = rewrite_block(chapter5, doc, {
        "第五章 核心技术路线与创新点": ("第五章 核心算法与创新点", 1),
        "5.1 技术路线": ("5.1 算法链路与三项核心创新", 2),
        "5.2 创新点一：多源异构数据交叉验证与时滞抑制": ("5.2 创新点一：多源岗位事实交叉验证与时滞抑制", 2),
        "5.3 创新点二：RAG约束与图谱事实校验的幻觉防控": ("5.3 创新点二：Graph-RAG双层校验与全链路溯源", 2),
        "5.4 创新点三：时序技能共现驱动的新岗位发现": ("5.4 创新点三：时序技能共现聚类与新岗位发现", 2),
        "5.5 创新点四：确定性评分与生成式解释解耦": ("5.5 支撑算法：六维确定性匹配与成长路径", 2),
    })
    app3 = rewrite_block(appendices["附录三 算法验证与可解释性说明"], doc, {
        "附录三 算法验证与可解释性说明": ("5.6 参数定义、判定门槛与可解释输出", 2),
        "A3.1 多源可信度加权": ("5.6.1 多源可信度加权", 3),
        "A3.2 时间衰减与能力演化": ("5.6.2 时间衰减与能力演化", 3),
        "A3.3 确定性匹配评分": ("5.6.3 确定性匹配评分", 3),
    })
    app11 = rewrite_block(appendices["附录十一 新岗位发现模型与判定规则"], doc, {
        "附录十一 新岗位发现模型与判定规则": ("5.7 新岗位候选指数与人工判定门槛", 2),
    }, strip_figures=True)
    app14 = rewrite_block(appendices["附录十四 六维确定性人岗匹配模型"], doc, {
        "附录十四 六维确定性人岗匹配模型": ("5.8 六维匹配公式与解释输出", 2),
    })

    chapter6 = rewrite_block(chapter6, doc, {
        "第六章 系统实现与演示": ("第六章 系统实现", 1),
        "6.1 当前数据与工程基线": ("6.1 开发框架与工程基线", 2),
        "6.2 演示实现证据": ("6.2 核心功能与运行效果", 2),
        "6.3 演示账号说明": ("6.3 源码结构与第三方依赖", 2),
    })
    app19 = rewrite_block(appendices["附录十九 原型系统页面与可演示证据"], doc, {
        "附录十九 原型系统页面与可演示证据": ("6.4 页面实现与功能证据", 2),
    }, strip_figures=True)
    app20 = rewrite_block(appendices["附录二十 部署拓扑、运行方式与离线能力"], doc, {
        "附录二十 部署拓扑、运行方式与离线能力": ("6.5 容器化部署、运行方式与离线能力", 2),
    }, strip_figures=True)

    chapter7 = rewrite_block(chapter7, doc, {"第七章 测试验证与效果评估": ("第七章 测试与评估", 1)})
    app21 = rewrite_block(appendices["附录二十一 测试策略、数据基线与可复现性"], doc, {
        "附录二十一 测试策略、数据基线与可复现性": ("7.4 测试策略、120条JD数据基线与可复现性", 2),
    }, strip_figures=True)
    app22 = rewrite_block(appendices["附录二十二 离线评测方法、指标定义与结果"], doc, {
        "附录二十二 离线评测方法、指标定义与结果": ("7.5 离线评测方法、指标定义与结果", 2),
    }, strip_figures=True)
    app23 = rewrite_block(appendices["附录二十三 典型错误、原因定位与改进实验"], doc, {
        "附录二十三 典型错误、原因定位与改进实验": ("7.6 典型错误、原因定位与改进实验", 2),
    })

    chapter8 = rewrite_block(chapter8, doc, {"第八章 应用场景、竞争分析与转化前景": ("第八章 总结与展望", 1)})
    closing_specs = [
        ("附录二十五 应用场景、试点设计与验收指标", "8.5 试点设计与验收指标", False),
        ("附录二十六 社会价值与公共利益边界", "8.6 社会价值与公共利益边界", False),
        ("附录二十七 成果转化模式与成本结构", "8.7 成果转化模式与成本结构", False),
        ("附录二十八 技术成熟度、应用成熟度与推进路线", "8.8 成熟度与推进路线", True),
        ("附录三十 风险矩阵与应对预案", "8.9 风险矩阵与应对预案", False),
    ]
    closing = []
    for source_title, new_title, strip in closing_specs:
        closing.extend(rewrite_block(appendices[source_title], doc, {source_title: (new_title, 2)}, strip_figures=strip))

    generated_requirements = generated(doc, lambda d: [d.element.body.append(e) for e in add_requirement_matrix(d)])
    generated_architecture = generated(doc, lambda d: [d.element.body.append(e) for e in add_figure(
        d,
        "image2科研重绘_图2_系统总体技术架构.png",
        "图10 数融智联系统总体技术架构与模块依赖",
    )])
    generated_implementation = generated(doc, lambda d: [d.element.body.append(e) for e in add_implementation_inventory(d)])
    generated_test_matrix = generated(doc, lambda d: [d.element.body.append(e) for e in add_test_case_matrix(d)])
    generated_ablation = generated(doc, lambda d: [d.element.body.append(e) for e in add_ablation_plan(d)])
    generated_conclusion = generated(doc, lambda d: [d.element.body.append(e) for e in add_conclusion(d)])
    generated_links = generated(doc, add_link_directory)

    assembly = []
    assembly.extend(front)
    assembly.extend(chapter1)
    assembly.extend(goal)
    assembly.extend(boundary)
    assembly.extend(solution)
    assembly.append(page_break_element(doc))
    assembly.append(add_heading(doc, "第二章 需求分析", 1))
    assembly.extend(users)
    assembly.extend(app1)
    assembly.extend(app8)
    assembly.extend(app24)
    assembly.extend(generated_requirements)
    assembly.append(page_break_element(doc))
    assembly.extend(chapter3)
    assembly.extend(generated_architecture)
    assembly.extend(app17)
    assembly.extend(chapter4_parts)
    assembly.append(page_break_element(doc))
    assembly.extend(chapter5)
    assembly.extend(app3)
    assembly.extend(app11)
    assembly.append(page_break_element(doc))
    assembly.extend(chapter6)
    assembly.extend(app19)
    assembly.extend(app20)
    assembly.extend(generated_implementation)
    assembly.append(page_break_element(doc))
    assembly.extend(chapter7)
    assembly.extend(app21)
    assembly.extend(app22)
    assembly.extend(app23)
    assembly.extend(generated_test_matrix)
    assembly.extend(generated_ablation)
    assembly.append(page_break_element(doc))
    assembly.extend(chapter8)
    assembly.extend(closing)
    assembly.extend(generated_conclusion)
    assembly.append(page_break_element(doc))
    assembly.extend(references)
    assembly.extend(generated_links)

    body = doc.element.body
    sect_pr = body.sectPr
    for element in list(body):
        if element is not sect_pr:
            body.remove(element)
    for element in assembly:
        if element.getparent() is not None:
            element.getparent().remove(element)
        body.insert(len(body) - 1, element)

    format_front(doc)
    normalise_document(doc)
    figures, captions = renumber_captions(doc)

    doc.core_properties.title = "数融智联：多源异构数据驱动的岗位能力图谱构建与动态演化分析系统"
    doc.core_properties.subject = "挑战杯项目申报书（会议优化版）"
    doc.core_properties.author = "数融智联项目团队"
    doc.save(OUTPUT)
    print(f"saved={OUTPUT}")
    print(f"tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)} figure_captions={figures} table_captions={captions}")
    print(f"cover={COVER}")


if __name__ == "__main__":
    main()
